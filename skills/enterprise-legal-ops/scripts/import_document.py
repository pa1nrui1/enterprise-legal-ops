#!/usr/bin/env python3
"""Import a document into the temporary intake area and extract text when possible."""

from __future__ import annotations

import argparse
import csv
import shutil
import subprocess
from datetime import datetime
from pathlib import Path
from uuid import uuid4


SUPPORTED = {".txt", ".md", ".csv", ".xlsx", ".xls", ".docx", ".doc", ".pdf", ".jpg", ".jpeg", ".png"}


def read_text(path: Path) -> tuple[str, str, str]:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".csv"}:
        return path.read_text(encoding="utf-8", errors="replace"), "native-text", ""
    if suffix == ".docx":
        try:
            from docx import Document  # type: ignore
        except Exception as exc:  # pragma: no cover - depends on optional env
            return "", "docx-unavailable", f"python-docx unavailable: {exc}"
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells]
                if any(values):
                    parts.append("\t".join(values))
        text = "\n".join(parts)
        return text, "python-docx", ""
    if suffix == ".doc":
        if not shutil.which("textutil"):
            return "", "doc-conversion-required", "legacy .doc requires conversion; textutil is unavailable"
        proc = subprocess.run(
            ["textutil", "-convert", "txt", "-stdout", str(path)],
            text=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
        )
        if proc.returncode != 0:
            return "", "doc-conversion-failed", proc.stderr.strip()
        return proc.stdout, "textutil", ""
    if suffix in {".xlsx", ".xls"}:
        if suffix == ".xls":
            try:
                import xlrd  # type: ignore
            except Exception:
                return "", "xls-unavailable", "legacy .xls requires xlrd; convert to .xlsx or install xlrd"
            book = xlrd.open_workbook(str(path))
            parts = []
            for sheet in book.sheets():
                parts.append(f"# Sheet: {sheet.name}")
                for r in range(sheet.nrows):
                    values = [str(sheet.cell_value(r, c)).strip() for c in range(sheet.ncols)]
                    if any(values):
                        parts.append("\t".join(values))
            return "\n".join(parts), "xlrd", ""
        try:
            import openpyxl  # type: ignore
        except Exception as exc:  # pragma: no cover
            return "", "xlsx-unavailable", f"openpyxl unavailable: {exc}"
        wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
        parts = []
        for ws in wb.worksheets:
            parts.append(f"# Sheet: {ws.title}")
            for row in ws.iter_rows(values_only=True):
                values = ["" if v is None else str(v) for v in row]
                if any(values):
                    parts.append("\t".join(values))
        return "\n".join(parts), "openpyxl", ""
    if suffix == ".pdf":
        try:
            import pypdf  # type: ignore
        except Exception as exc:  # pragma: no cover
            return "", "pdf-unavailable", f"pypdf unavailable: {exc}"
        reader = pypdf.PdfReader(str(path))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        return text, "pypdf", ""
    if suffix in {".jpg", ".jpeg", ".png"}:
        if not shutil.which("tesseract"):
            return "", "ocr-required", "image import recorded; OCR is required before substantive use"
        proc = subprocess.run(
            ["tesseract", str(path), "stdout", "-l", "chi_sim+eng"],
            text=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
        )
        if proc.returncode != 0:
            return "", "ocr-failed", proc.stderr.strip()
        return proc.stdout, "tesseract", ""
    return "", "unsupported", f"unsupported suffix: {suffix}"


def extract_key_fields(text: str) -> str:
    if not text:
        return ""
    fields = []
    dates = sorted(set(__import__("re").findall(r"\d{4}[-年./]\d{1,2}[-月./]\d{1,2}日?", text)))[:10]
    amounts = sorted(set(__import__("re").findall(r"(?:人民币)?\s*\d+(?:,\d{3})*(?:\.\d+)?\s*(?:元|万元|人民币)?", text)))[:10]
    ids = sorted(set(__import__("re").findall(r"\b\d{17}[\dXx]\b", text)))[:5]
    if dates:
        fields.append("日期=" + "、".join(dates))
    if amounts:
        fields.append("金额=" + "、".join(amounts))
    if ids:
        fields.append("身份证号=" + "、".join(ids))
    return "；".join(fields)


def write_read_review(workspace: Path, import_id: str, source: Path, method: str, status: str, key_fields: str, error: str, text: str) -> Path:
    out_dir = workspace / "_system" / "read-reviews"
    out_dir.mkdir(parents=True, exist_ok=True)
    out = out_dir / f"{import_id}-read-review-summary.md"
    completeness = "可用" if text else "不可用/需补充处理"
    doubts = error or ("无明显存疑项" if text else "未取得可用文本")
    out.write_text(
        "\n".join([
            "# 读取复查摘要",
            "",
            "## 文件",
            f"- {source}",
            "",
            "## 读取方式",
            f"- {method}",
            "",
            "## 关键字段",
            f"- {key_fields or '未自动识别'}",
            "",
            "## 存疑项",
            f"- {doubts}",
            "",
            "## 完整性评估",
            f"- {completeness}",
            "",
            "## 是否需要用户确认",
            f"- {'是' if error or not text else '否'}",
            "",
        ]),
        encoding="utf-8",
    )
    log = workspace / "_system" / "read-review-log.md"
    log.parent.mkdir(parents=True, exist_ok=True)
    with log.open("a", encoding="utf-8") as f:
        f.write(f"\n## {import_id}\n")
        f.write(f"- 文件：{source}\n")
        f.write(f"- 读取方式：{method}\n")
        f.write(f"- 读取状态：{status}\n")
        f.write(f"- 读取复查摘要：{out}\n")
    return out


def write_source_boundary(workspace: Path, import_id: str, source: Path, status: str, error: str) -> Path:
    out_dir = workspace / "_system" / "source-boundaries"
    out_dir.mkdir(parents=True, exist_ok=True)
    out = out_dir / f"{import_id}-source-boundary.md"
    out.write_text(
        "\n".join([
            "# 来源边界",
            "",
            "## 已读取材料",
            f"- {source}" if status == "成功" else "- 无完整可用读取材料",
            "",
            "## 已核验内容",
            "- 未进行法规或规则核验。",
            "",
            "## 未读取或缺失材料",
            f"- {source}：{error}" if status != "成功" else "- 无",
            "",
            "## 存疑项",
            f"- {error or '无明显存疑项'}",
            "",
            "## 输出边界",
            "- 本次仅完成文件导入和读取记录，不构成法律判断或正式法律意见。",
            "",
        ]),
        encoding="utf-8",
    )
    log = workspace / "_system" / "source-boundary-log.md"
    log.parent.mkdir(parents=True, exist_ok=True)
    with log.open("a", encoding="utf-8") as f:
        f.write(f"\n## {import_id}\n")
        f.write(f"- 来源边界记录：{out}\n")
    return out


def append_import_log(workspace: Path, row: dict[str, str]) -> None:
    path = workspace / "_system" / "import-log.csv"
    path.parent.mkdir(parents=True, exist_ok=True)
    exists = path.exists()
    headers = ["import_id", "file_path", "module", "status", "notes", "created_at"]
    with path.open("a", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=headers)
        if not exists:
            writer.writeheader()
        writer.writerow(row)


def append_extracted_index(workspace: Path, row: dict[str, str]) -> None:
    path = workspace / "05-本地问库" / "extracted-text-index.csv"
    path.parent.mkdir(parents=True, exist_ok=True)
    exists = path.exists()
    headers = [
        "text_id", "module", "object_id", "source_file_path", "extracted_text_path",
        "read_method", "read_status", "ocr_status", "key_fields", "created_at", "updated_at",
    ]
    with path.open("a", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=headers)
        if not exists:
            writer.writeheader()
        writer.writerow(row)


def main() -> int:
    parser = argparse.ArgumentParser(description="Import a document into an Enterprise Legal Ops workspace.")
    parser.add_argument("--workspace", required=True)
    parser.add_argument("--file", required=True)
    parser.add_argument("--module", default="待识别")
    parser.add_argument("--object-id", default="")
    args = parser.parse_args()

    workspace = Path(args.workspace).expanduser().resolve()
    src = Path(args.file).expanduser().resolve()
    if not src.exists():
        raise SystemExit(f"file not found: {src}")
    if src.suffix.lower() not in SUPPORTED:
        raise SystemExit(f"unsupported file type: {src.suffix}")

    now = datetime.now().isoformat(timespec="seconds")
    import_id = f"IMP-{uuid4().hex[:8]}"
    intake_dir = workspace / "06-导入暂存" / "待识别"
    intake_dir.mkdir(parents=True, exist_ok=True)
    dest = intake_dir / f"{import_id}-{src.name}"
    shutil.copy2(src, dest)

    text, method, error = read_text(dest)
    read_status = "成功" if text else "需处理"
    if error and method not in {"ocr-required"}:
        read_status = "失败"
    extracted_path = ""
    if text:
        extracted_dir = workspace / "05-本地问库" / "读取文本"
        extracted_dir.mkdir(parents=True, exist_ok=True)
        extracted = extracted_dir / f"{import_id}.txt"
        extracted.write_text(text, encoding="utf-8")
        extracted_path = str(extracted)
    key_fields = extract_key_fields(text)
    read_review_path = write_read_review(workspace, import_id, dest, method, read_status, key_fields, error, text)
    source_boundary_path = write_source_boundary(workspace, import_id, dest, read_status, error)

    append_import_log(workspace, {
        "import_id": import_id,
        "file_path": str(dest),
        "module": args.module,
        "status": read_status,
        "notes": error,
        "created_at": now,
    })
    append_extracted_index(workspace, {
        "text_id": f"TXT-{uuid4().hex[:8]}",
        "module": args.module,
        "object_id": args.object_id,
        "source_file_path": str(dest),
        "extracted_text_path": extracted_path,
        "read_method": method,
        "read_status": read_status,
        "ocr_status": "待OCR" if method == "ocr-required" else "",
        "key_fields": key_fields,
        "created_at": now,
        "updated_at": now,
    })

    print("文件已导入。")
    print(f"导入编号：{import_id}")
    print(f"暂存位置：{dest}")
    print(f"读取方式：{method}")
    print(f"读取状态：{read_status}")
    if extracted_path:
        print(f"读取文本：{extracted_path}")
    print(f"读取复查摘要：{read_review_path}")
    print(f"来源边界记录：{source_boundary_path}")
    if error:
        print(f"存疑/错误：{error}")
    print("下一步：请确认文件所属模块和关键字段；涉及法律判断前必须完成读取复查。")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
